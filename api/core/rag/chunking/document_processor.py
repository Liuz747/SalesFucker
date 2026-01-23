"""
文档处理器

该模块提供文档处理管道，将上传的文档转换为可检索的分块。
支持多种文件格式（TXT, MD, PDF, DOCX）的解析和处理。

核心功能:
- 文件格式解析（TXT, MD, PDF, DOCX）
- 文本提取和清理
- 文档分块
- 元数据提取
- 批量处理
"""

from pathlib import Path
from typing import Optional

from models import DocumentType
from utils import get_component_logger
from .text_splitter import TextSplitter

logger = get_component_logger(__name__, "DocumentProcessor")


class DocumentProcessor:
    """
    文档处理器

    处理各种格式的文档，提取文本并分块。
    """

    def __init__(self, text_splitter: Optional[TextSplitter] = None):
        """
        初始化DocumentProcessor

        参数:
            text_splitter: 文本分块器实例（可选）
        """
        self.text_splitter = text_splitter or TextSplitter()

    def process_document(
        self,
        file_path: str,
        file_type: DocumentType,
        metadata: Optional[dict] = None
    ) -> list[dict]:
        """
        处理文档并生成分块

        参数:
            file_path: 文件路径
            file_type: 文件类型
            metadata: 文档元数据

        返回:
            list[dict]: 分块列表
        """
        try:
            logger.info(f"处理文档: {file_path}, 类型: {file_type}")

            # 提取文本
            text = self._extract_text(file_path, file_type)

            if not text or not text.strip():
                logger.warning(f"文档为空或无法提取文本: {file_path}")
                return []

            # 清理文本
            text = self._clean_text(text)

            # 分块
            if file_type == DocumentType.MD:
                chunks = self.text_splitter.split_markdown(text, metadata)
            else:
                chunks = self.text_splitter.split_text(text, metadata)

            logger.info(f"文档处理完成: {len(chunks)} 个分块")
            return chunks

        except Exception as e:
            logger.error(f"处理文档失败: {file_path}, 错误: {e}")
            raise

    def _extract_text(self, file_path: str, file_type: DocumentType) -> str:
        """
        从文件提取文本

        参数:
            file_path: 文件路径
            file_type: 文件类型

        返回:
            str: 提取的文本
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        match file_type:
            case DocumentType.TXT | DocumentType.MD:
                return self._extract_text_from_txt(path)
            case DocumentType.PDF:
                return self._extract_text_from_pdf(path)
            case DocumentType.DOCX:
                return self._extract_text_from_docx(path)
            case _:
                logger.warning(f"不支持的文件类型: {file_type}")
                return ""

    def _extract_text_from_txt(self, path: Path) -> str:
        """
        从TXT/MD文件提取文本

        参数:
            path: 文件路径

        返回:
            str: 文本内容
        """
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.debug(f"使用编码 {encoding} 读取文件成功")
                    return text
                except UnicodeDecodeError:
                    continue

            # 如果所有编码都失败，使用二进制模式并忽略错误
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            logger.warning(f"使用UTF-8（忽略错误）读取文件")
            return text

        except Exception as e:
            logger.error(f"读取TXT文件失败: {e}")
            raise

    def _extract_text_from_pdf(self, path: Path) -> str:
        """
        从PDF文件提取文本

        参数:
            path: 文件路径

        返回:
            str: 文本内容
        """
        try:
            import pypdf

            text_parts = []

            with open(path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"提取PDF第{page_num + 1}页失败: {e}")

            text = '\n\n'.join(text_parts)
            logger.info(f"PDF提取完成: {len(pdf_reader.pages)} 页")
            return text

        except ImportError:
            logger.error("pypdf未安装，无法处理PDF文件")
            raise ImportError("请安装pypdf: pip install pypdf")
        except Exception as e:
            logger.error(f"读取PDF文件失败: {e}")
            raise

    def _extract_text_from_docx(self, path: Path) -> str:
        """
        从DOCX文件提取文本

        参数:
            path: 文件路径

        返回:
            str: 文本内容
        """
        try:
            import docx

            doc = docx.Document(path)
            text_parts = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            text = '\n\n'.join(text_parts)
            logger.info(f"DOCX提取完成: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格")
            return text

        except ImportError:
            logger.error("python-docx未安装，无法处理DOCX文件")
            raise ImportError("请安装python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"读取DOCX文件失败: {e}")
            raise

    def _clean_text(self, text: str) -> str:
        """
        清理文本

        参数:
            text: 原始文本

        返回:
            str: 清理后的文本
        """
        # 移除多余的空白字符
        text = '\n'.join(line.strip() for line in text.split('\n'))

        # 移除多余的空行（保留最多2个连续空行）
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 移除特殊控制字符
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

        return text.strip()

    def process_batch(
        self,
        documents: list[tuple[str, DocumentType, dict]]
    ) -> dict[str, list[dict]]:
        """
        批量处理文档

        参数:
            documents: (file_path, file_type, metadata) 列表

        返回:
            dict[str, list[dict]]: {file_path: chunks} 映射
        """
        results = {}

        for file_path, file_type, metadata in documents:
            try:
                chunks = self.process_document(file_path, file_type, metadata)
                results[file_path] = chunks
            except Exception as e:
                logger.error(f"批量处理失败: {file_path}, 错误: {e}")
                results[file_path] = []

        logger.info(f"批量处理完成: {len(results)} 个文档")
        return results


# 全局DocumentProcessor实例
document_processor = DocumentProcessor()
